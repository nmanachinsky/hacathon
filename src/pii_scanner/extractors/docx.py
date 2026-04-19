"""DOCX → текст."""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from ..types import TextChunk


def extract(path: Path) -> Iterable[TextChunk]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)
    if parts:
        yield TextChunk(text="\n".join(parts), locator="docx_body")
